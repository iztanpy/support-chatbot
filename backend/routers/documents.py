from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from database import get_conn, get_live_version_id
import anthropic
import os
import json
import uuid
import io

router = APIRouter()
client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY", ""))

# ── text extraction ───────────────────────────────────────────────────────────

def extract_text_from_pdf(data: bytes) -> str:
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(data))
        parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
        return "\n\n".join(parts)
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"PDF extraction failed: {e}")

def extract_text_from_docx(data: bytes) -> str:
    try:
        import docx
        doc = docx.Document(io.BytesIO(data))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"DOCX extraction failed: {e}")

def extract_text(filename: str, data: bytes) -> str:
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(data)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(data)
    elif ext in ("txt", "md", "markdown"):
        return data.decode("utf-8", errors="replace")
    else:
        raise HTTPException(status_code=422, detail=f"Unsupported file type: .{ext}")

# ── Claude parse ──────────────────────────────────────────────────────────────

PARSE_SYSTEM = """You are a configuration extraction assistant for an AI customer service bot.

Given a document, extract any of the following if present:
1. Bot guidelines / behaviour rules
2. Knowledge base URL
3. Workflow definitions (trigger phrases + API endpoints)
4. Known Q&A corrections (question + correct answer pairs)

Return ONLY a valid JSON object with this exact structure (omit keys you have no data for):
{
  "guidelines": "extracted guidelines text or null",
  "knowledge_base_url": "URL string or null",
  "workflows": [
    {
      "trigger": "phrase that triggers this workflow",
      "description": "what it does",
      "endpoint_url": "https://...",
      "http_method": "GET or POST",
      "inputs": [
        {"name": "param_name", "label": "human prompt", "required": true, "param_type": "query"}
      ],
      "response_template": "template with {{field}} placeholders"
    }
  ],
  "corrections": [
    {"question_pattern": "question text", "correct_answer": "correct answer text"}
  ],
  "confidence_notes": "brief note on extraction quality or missing/ambiguous data"
}

Be conservative — only extract data that is clearly present. Do not invent endpoints or guidelines."""

async def parse_document_with_claude(text: str, filename: str) -> dict:
    prompt = f"""Document filename: {filename}

Document content:
---
{text[:12000]}
---

Extract configuration as instructed."""

    try:
        response = client.messages.create(
            model="claude-opus-4-5",
            max_tokens=4096,
            system=PARSE_SYSTEM,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = response.content[0].text.strip()
        # Strip markdown fences if present
        if raw.startswith("```"):
            raw = raw.split("\n", 1)[1].rsplit("```", 1)[0]
        return json.loads(raw)
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=422, detail=f"Could not parse AI response as JSON: {e}")
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"AI parsing failed: {e}")

# ── diff helper ───────────────────────────────────────────────────────────────

def build_diff(current: dict, proposed: dict) -> dict:
    """
    Build a simple before/after diff for the preview UI.
    Only includes fields that actually changed.
    """
    diff = {}

    if proposed.get("guidelines") and proposed["guidelines"] != current.get("guidelines"):
        diff["guidelines"] = {
            "before": current.get("guidelines", ""),
            "after": proposed["guidelines"],
        }

    if proposed.get("knowledge_base_url") and proposed["knowledge_base_url"] != current.get("knowledge_base_url"):
        diff["knowledge_base_url"] = {
            "before": current.get("knowledge_base_url", ""),
            "after": proposed["knowledge_base_url"],
        }

    if proposed.get("workflows"):
        diff["workflows_to_add"] = proposed["workflows"]

    if proposed.get("corrections"):
        diff["corrections_to_add"] = proposed["corrections"]

    return diff

# ── endpoints ─────────────────────────────────────────────────────────────────

@router.post("/parse")
async def parse_document(file: UploadFile = File(...)):
    """
    Step 1: Upload and parse a document.
    Returns extracted config + a diff against the current live version.
    Does NOT apply anything — purely a preview.
    """
    data = await file.read()
    text = extract_text(file.filename, data)

    if not text.strip():
        raise HTTPException(status_code=422, detail="Document appears to be empty or unreadable")

    proposed = await parse_document_with_claude(text, file.filename)

    # Load current live config for diff
    version_id = get_live_version_id()
    conn = get_conn()
    config_row = conn.execute(
        "SELECT * FROM bot_config WHERE version_id=?", (version_id,)
    ).fetchone()
    conn.close()
    current = dict(config_row) if config_row else {}

    diff = build_diff(current, proposed)

    return {
        "filename": file.filename,
        "extracted": proposed,
        "diff": diff,
        "has_changes": bool(diff),
        "confidence_notes": proposed.get("confidence_notes", ""),
    }

@router.post("/apply")
async def apply_document(body: dict):
    """
    Step 2: Apply a parsed document result as a new draft version.
    Expects: { version_name, extracted: {...}, clone_from_version_id?, source_document? }
    Creates a new draft (not live) — user must promote separately.
    """
    extracted       = body.get("extracted", {})
    version_name    = body.get("version_name", "Untitled Upload")
    source_document = body.get("source_document", "")

    # Document is the sole source of truth — nothing is inherited from the live version.
    # If the document omits a field, that field is blank/empty in the new version.
    conn = get_conn()

    # Create new version row
    conn.execute(
        "INSERT INTO bot_versions (name, is_live, source_document) VALUES (?, 0, ?)",
        (version_name, source_document)
    )
    conn.commit()
    new_version_id = conn.execute(
        "SELECT id FROM bot_versions ORDER BY id DESC LIMIT 1"
    ).fetchone()["id"]

    # Config — document only, no fallback
    conn.execute("""
        INSERT INTO bot_config (version_id, knowledge_base_url, guidelines)
        VALUES (?, ?, ?)
    """, (
        new_version_id,
        extracted.get("knowledge_base_url") or "",
        extracted.get("guidelines") or "",
    ))

    if extracted.get("workflows"):
        # Workflows — document only, no fallback
        for wf in extracted.get("workflows", []):
            conn.execute("""
                INSERT INTO workflows
                  (id, version_id, trigger, description, endpoint_url, http_method,
                   inputs, headers, response_template, enabled)
            VALUES (?,?,?,?,?,?,?,?,?,?)
        """, (
            f"w{uuid.uuid4().hex[:8]}", new_version_id,
            wf.get("trigger", ""), wf.get("description", ""),
            wf.get("endpoint_url", ""), wf.get("http_method", "GET"),
            json.dumps(wf.get("inputs", [])), json.dumps(wf.get("headers", {})),
            wf.get("response_template", ""), 1,
        ))

    # Corrections — document only, no fallback
    if extracted.get("corrections"):
        for corr in extracted.get("corrections", []):
            conn.execute("""
                INSERT INTO corrections (version_id, question_pattern, correct_answer)
                VALUES (?,?,?)
            """, (new_version_id, corr.get("question_pattern", ""), corr.get("correct_answer", "")))

    conn.commit()
    row = conn.execute("SELECT * FROM bot_versions WHERE id=?", (new_version_id,)).fetchone()
    conn.close()

    return {
        "message": f"Draft version '{version_name}' created successfully.",
        "version": {
            "id": row["id"],
            "name": row["name"],
            "is_live": bool(row["is_live"]),
            "source_document": row["source_document"],
            "created_at": row["created_at"],
        }
    }
